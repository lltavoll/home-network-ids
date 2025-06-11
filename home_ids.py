from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Define a helper function for formatting headings
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

# Define a helper function to add a paragraph with optional bold text
def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p

# Header
add_paragraph("Gustavo Maldonado", bold=True)
add_paragraph("lltavolll@gmail.com • (602)-451-5832 • Peoria, AZ")

# Summary
add_heading("SUMMARY", level=2)
add_paragraph("Cybersecurity Intern | Arizona Department of Homeland Security (Application) Cybersecurity student at Bellevue University (B.S. expected Winter 2026) with 5+ years of experience in secure asset management, risk mitigation, and compliance. Skilled in network security, threat analysis, and operational security, with hands-on experience in NIST-based compliance frameworks, security assessments, and risk mitigation. Eager to contribute to public sector cybersecurity efforts while further developing expertise in threat detection, malware analysis, and vulnerability scanning.")

# Experience
add_heading("EXPERIENCE", level=2)

add_paragraph("Target Distribution Center, Warehouse Worker | Operational Risk & Asset Control", bold=True)
add_paragraph("Aug 2022 – Present")
add_paragraph("• Conduct rigorous quality control checks to ensure accuracy and integrity of outbound shipments, mitigating risks of data and asset loss.")
add_paragraph("• Maintain secure and accurate inventory records, retrieve and track high-value items, and resolve discrepancies quickly.")
add_paragraph("• Follow strict safety and compliance protocols, ensuring secure handling and transport of company assets.")
add_paragraph("• Operate and manage digital warehouse systems (WMS), reinforcing attention to systems security and process accuracy.")

add_paragraph("Dish Network, Inventory Specialist (ISP) | Asset Management & Systems Coordination", bold=True)
add_paragraph("Jul 2018 – Mar 2020")
add_paragraph("• Managed high-volume inventory flow, ensuring secure storage, accurate tracking, and timely distribution of critical materials and equipment.")
add_paragraph("• Oversaw processing of vendor shipments, returns, and customer deliveries, ensuring data accuracy and process compliance.")
add_paragraph("• Coordinated Samsung fulfillment, tracking job statuses and ETAs—skills transferable to monitoring system statuses and threat timelines.")
add_paragraph("• Maintained secure documentation and reporting practices to safeguard inventory data integrity.")

add_paragraph("Liberty National, Licensed Insurance Agent | Risk Assessment & Client Security", bold=True)
add_paragraph("Apr 2020 – Jul 2021")
add_paragraph("• Conducted in-depth risk assessments for clients, which sharpened principles of threat modeling and mitigation planning.")
add_paragraph("• Handled sensitive customer information, ensuring data protection and compliance with industry privacy standards.")
add_paragraph("• Utilized CRM tools for secure information storage and efficient follow-ups, developing organized, secure documentation skills.")
add_paragraph("• Communicated complex policy options clearly — a key transferable skill for stakeholder communication in cybersecurity.")

add_paragraph("AZ Tech, Front Desk/Sales Associate | Technical Support & Secure Data Handling", bold=True)
add_paragraph("Mar 2015 – Dec 2017")
add_paragraph("• Managed repair intake and provided technical advice on device security and recovery solutions.")
add_paragraph("• Recorded and stored customer data securely using Excel spreadsheets, maintaining privacy protocols.")
add_paragraph("• Assisted in minor malware cleanups and antivirus installations, introducing customers to best practices in digital hygiene.")
add_paragraph("• Upsold security-enhancing products like VPN-enabled routers and encrypted drives.")

add_paragraph("United Parcel Service (UPS), Package Handler", bold=True)
add_paragraph("Jan 2014 – Feb 2015")
add_paragraph("• Loaded packages into trucks, ensuring accurate and efficient handling.")
add_paragraph("• Utilized hand-held scanners and maintained logs to accurately track item movements and inventory.")
add_paragraph("• Resolved conveyor jams and blockages to prevent delays.")

# Education
add_heading("EDUCATION", level=2)
add_paragraph("Bellevue University", bold=True)
add_paragraph("Bachelor of Science • Cybersecurity")
add_paragraph("Expected Graduation: Winter 2026")

# Projects
add_heading("PROJECTS", level=2)
add_paragraph("Home Network Threat Detection & Alert System", bold=True)
add_paragraph("Personal Project | Python, Wireshark, Email Automation, Windows Firewall Logs")
add_paragraph("• Designed and implemented a lightweight intrusion detection system (IDS) for a home network using Python and Windows firewall logs.")
add_paragraph("• Parsed logs for suspicious IPs and behaviors using regular expressions and automated alerts via email when anomalies were detected.")
add_paragraph("• Integrated open-source threat feeds to enrich detection logic with blacklisted IPs.")
add_paragraph("• Used Wireshark to verify alerts and understand packet-level traffic patterns.")

# Skills
add_heading("SKILLS", level=2)
add_paragraph("Network Security & Access Control • Cyber Threat Analysis • Database Security • Secure Asset Management • Risk Mitigation • Incident Resolution • Operating Systems (Windows/Linux) • Programming Fundamentals (Python, etc.) • MS Office (Excel, Word, Outlook) • Oracle")

# Awards
add_heading("HONORS & AWARDS", level=2)
add_paragraph("Dean's List")
add_paragraph("Bellevue University • May 2025")
add_paragraph("Dean's List 4.0 GPA")

# Save the document
doc_path = "/mnt/data/Updated_Resume_Gustavo_Maldonado.docx"
doc.save(doc_path)

doc_path
